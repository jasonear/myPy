from docx import Document

doc = Document('wrfile\\test-table.docx')

# # 遍历段落
# for i in range(0,15):
#     para = doc.paragraphs[i]
#     print("第"+str(i)+"段："+para.text)

#     # 遍历段落中的run
#     for j in range(0,len(para.runs)):
#         myrun = para.runs[j]
#         print( str(j) +':' +myrun.text)



# iterator = iter(doc.element.body)

#     # 遍历文档的每个元素
# for element in iterator:
#     # print(":"+element.text)

#     # # 检查元素是否包含特定的文字
#     if element.text is not None and "安全宜居" in element.text:
#         print(element.text)

#         # 如果找到了特定的文字，那么继续遍历直到找到表格
#         for element in iterator:
#             if element.tag.endswith('tbl'):
#                 table = Table(element, doc)
#                 # 打印表格的内容
#                 for row in table.rows:
#                     for cell in row.cells:
#                         print(cell.text)
#             # 如果遇到另一个文字元素，就停止打印表格
#             elif element.tag.endswith('p'):
#                 break


# found_text = False
# # 遍历文档的每个段落
# for para in doc.paragraphs:
#     print(":" + para.text)
#     # 检查段落是否包含特定的文字
#     if "安全宜居" in para.text:
#         found_text = True
#         print("ok------------------")

#     # 如果找到了特定的文字，那么打印后续的所有表格
#     if found_text and 'tbl' in para._p.xml:
#         print("table------------------")
#         table = [cell.text for row in para._p.table.rows for cell in row.cells]
#         print(table)


# #OK----------------------
# print_flag = False

# for element in doc.element.body:
#     if element.tag.endswith('p'):
#         para = doc._element.xpath(element.getroottree().getpath(element))[0]
#         if "安全宜居" in para.text:
#             print_flag = True
#         elif print_flag:
#             # 如果遇到另一个文字元素，就停止打印表格
#             break
#     elif print_flag and element.tag.endswith('tbl'):
#         table = doc._element.xpath(element.getroottree().getpath(element))[0]
#         for row in table.xpath('.//w:tr'):
#             cells = row.xpath('.//w:tc')
#             for cell in cells:
#                 print(''.join(node.text for node in cell.xpath('.//w:t')))



print_flag = False
mytbs=[]
mystr="安全宜居"

for element in doc.element.body:
    if element.tag.endswith('p'):
        para = doc._element.xpath(element.getroottree().getpath(element))[0]
        if mystr in para.text:
            print_flag = True
        elif print_flag and mystr not in para.text:
            # 如果遇到另一个不包含“安全宜居”的文字元素，就停止打印表格
            print_flag = False
    elif print_flag and element.tag.endswith('tbl'):
        table = doc._element.xpath(element.getroottree().getpath(element))[0]
        myrows=[]

        for row in table.xpath('.//w:tr'):
            cells = row.xpath('.//w:tc')
            mycells=[]

            for cell in cells:
                mycells.append(''.join(node.text for node in cell.xpath('.//w:t')))
                # print(''.join(node.text for node in cell.xpath('.//w:t')))
            print(mycells)
            myrows.append(mycells)

        print("\n")
        mytbs.append(myrows)

print(mytbs)