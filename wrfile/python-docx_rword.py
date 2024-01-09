from docx import Document

# 打开Word文档
doc = Document('wrfile\example.docx')

# # 遍历文档中的所有段落
# for para in doc.paragraphs:
#     print(para.text)

# 遍历文档中的指定段落
for i in range(3,8):
    para = doc.paragraphs[i]
    print(para.text)


# # 遍历文档中的所有表格
# for table in doc.tables:  # 遍历每个表格
#     for row in table.rows: # 遍历表格的每一行
#         for cell in row.cells: # 遍历行里的每个单元格
#             print(cell.text)

# # 遍历文档中的所有图片
# for inline_shape in doc.inline_shapes:
#     if inline_shape.type == 3:
#         print(inline_shape.image.filename)

# # 修改文档中的内容
# doc.paragraphs[0].text = '这是修改后的内容'
